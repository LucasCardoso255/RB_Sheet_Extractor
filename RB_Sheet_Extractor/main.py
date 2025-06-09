import os
import docx
import re
import openpyxl
from openpyxl.styles import Font, Alignment

def extract_word_data(file_path):
    data = {
        "product_name": "N/A",
        "client": "N/A",
        "production_per_hour": "N/A",
        "channel_break": "Não",
        "deburring": "Não",
        "sandblasting": "Não",
        "sanding": "Não",
        "packaging_inspection": "Não",
        "packaging_type": "Não Informado",
        "quantity_per_package": "N/A",
    }

    try:
        doc = docx.Document(file_path)

        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + " "
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += cell.text + " "

        full_text = re.sub(r'\s+', ' ', full_text.replace("\xa0", " ").replace("\t", " ").strip())

        name_match = re.search(r"Nome:\s*(.*?)\s*Código:", full_text, re.DOTALL)
        client_match = re.search(r"Cliente:\s*([^\n\r]+?)\s", full_text)
        production_match = re.search(r"(\d+)\s*Peças", full_text, re.IGNORECASE)

        if name_match:
            data["product_name"] = name_match.group(1).strip()
        if client_match:
            data["client"] = client_match.group(1).strip()
        if production_match:
            data["production_per_hour"] = production_match.group(1).strip()

        packaging_text_raw = ""
        packaging_match = re.search(r"embalar as peças em\s*([^\.]+?(?:N°\s*\d+[A-Z])?(?:Nº\s*\d+[A-Z])?)\.", full_text, re.IGNORECASE)
        if packaging_match:
            packaging_text_raw = packaging_match.group(1).strip()
            packaging_text_raw = re.sub(r',\s*(?:e enviá-las para o cliente|separadas por modelo|modelos separados e envolver em plástico filme|e envolver em plástico filme).*$', '', packaging_text_raw, flags=re.IGNORECASE).strip()
            packaging_text_raw = re.sub(r'\s*Quant.*$', '', packaging_text_raw, flags=re.IGNORECASE).strip()
            packaging_text_raw = re.sub(r'[,\.]$', '', packaging_text_raw).strip()
            packaging_text_raw = packaging_text_raw.replace('Nº', 'N°')
            packaging_text_raw = re.sub(r'sacos$', 'saco', packaging_text_raw, flags=re.IGNORECASE).strip()

        if re.search(r'\b(?:caixa|caixas)\b', packaging_text_raw, re.IGNORECASE):
            box_number_match = re.search(r'N[°º]\s*(0?[1-9]|10|11)[A-Z]?', packaging_text_raw, re.IGNORECASE)
            if box_number_match:
                data["packaging_type"] = f"Caixa de Papelão N° {box_number_match.group(1).upper()}"
            else:
                data["packaging_type"] = "Caixa de Papelão"
        elif re.search(r'\b(?:pallet|palete)\b', packaging_text_raw, re.IGNORECASE):
            data["packaging_type"] = "Pallet"
        elif re.search(r'\b(?:caixa de pl[aá]stico|plastico)\b', packaging_text_raw, re.IGNORECASE):
            data["packaging_type"] = "Caixa de Plástico"
        elif re.search(r'\b(?:saco de r[aá]fia|saco)\b', packaging_text_raw, re.IGNORECASE):
            data["packaging_type"] = "Saco de Ráfia"
        elif packaging_text_raw:
            data["packaging_type"] = packaging_text_raw


        quantity_per_package_raw = "N/A"
        quantity_match = re.search(r"Quant(?:\.\s*| )Por\s*(?:saco|caixa|peça|caixas|sacos|peças)?:\s*(.*?)(?:\n|\Z)", full_text, re.IGNORECASE | re.DOTALL)
        if quantity_match:
            quantity_per_package_raw = quantity_match.group(1).strip()
        else:
            fallback_quantity_pattern = r"(?:(?:[^\d\n]*(?:Conforme Pedido do Cliente|TOLERÂNCIA DE PARÂMETRO NA 450TON|DESCRIÇÃO TOLERÂNCIA).*?)?\s*)*?"
            fallback_quantity_pattern += r"((?:[\w\s]+?:\s*\d+(?:\.\d+)?\s*peças\b(?:\s*por Caixa\b)?(?:\s*cada\b)?(?:;|\s)*)*"
            fallback_quantity_pattern += r"\d+(?:\.\d+)?\s*peças\b(?:\s*cada\b)?(?: por (?:Caixa|saco))?(?:;|\s)*)"
            
            fallback_match = re.search(fallback_quantity_pattern, full_text, re.IGNORECASE | re.DOTALL)
            if fallback_match:
                quantity_per_package_raw = fallback_match.group(1).strip()
                quantity_per_package_raw = re.sub(r'(?:Conforme Pedido do Cliente|TOLERÂNCIA DE PARÂMETRO NA 450TON|DESCRIÇÃO TOLERÂNCIA|Rampa Fase|Vel\. Multip|Atraso Mult|Partida Fase|Pressão Multip|Tempo Comp|Tempo Resf|Vel\. Acomp\. Molde|Parar Injeção|Atraso Retorno).*$', '', quantity_per_package_raw, flags=re.IGNORECASE | re.DOTALL).strip()
                quantity_per_package_raw = re.sub(r'[^\d\w:;\s\.]', '', quantity_per_package_raw).strip()

        if quantity_per_package_raw != "N/A":
            items = re.findall(r"([\w\s]+?:\s*\d+(?:\.\d+)?\s*peças\b(?:\s*por Caixa\b)?(?:\s*cada\b)?|\d+(?:\.\d+)?\s*peças\b(?:\s*cada\b)?)", quantity_per_package_raw, re.IGNORECASE)
            if items:
                formatted_quantities = []
                for item in items:
                    cleaned_item = re.sub(r'\s*(?:por Caixa|cada)\s*$', '', item.strip(), flags=re.IGNORECASE)
                    formatted_quantities.append(cleaned_item)
                data["quantity_per_package"] = "; ".join(formatted_quantities).strip()
                data["quantity_per_package"] = re.sub(r'peças peças', 'peças', data["quantity_per_package"], flags=re.IGNORECASE)
            else:
                data["quantity_per_package"] = "N/A"
        else:
            data["quantity_per_package"] = "N/A"

        if "INSPECAO VISUAL 100%" in full_text.upper() or "INSPEÇÃO FINAL" in full_text.upper():
            data["packaging_inspection"] = "Sim"
        if "QUEBRA DO CANAL" in full_text.upper():
            data["channel_break"] = "Sim"
        if "REBARBAÇÃO" in full_text.upper():
            data["deburring"] = "Sim"
        if "JATO DE GRANALHA" in full_text.upper():
            data["sandblasting"] = "Sim"
        if "LIXA" in full_text.upper():
            data["sanding"] = "Sim"

        return data

    except Exception as e:
        print(f"Error processing {file_path}: {e}")
        return None

def process_folders_and_create_excel(root_folder, output_excel_file):
    workbook = openpyxl.Workbook()
    sheet = workbook.active

    headers = [
        "Nome do Produto", "Cliente", "Producao/Hora", "Quebra de Canal",
        "Rebarbacao", "Jateamento", "Lixa", "Inspecao e Embalagem",
        "Tipo de Embalagem", "Quant. por Embalagem",
    ]

    sheet.append(headers)

    bold_font = Font(bold=True)
    center_align = Alignment(horizontal='center', vertical='center')
    for cell in sheet[1]:
        cell.font = bold_font
        cell.alignment = center_align

    sheet.column_dimensions['A'].width = 51
    for col_letter in ['B', 'C', 'D', 'E', 'F', 'G', 'H']:
        sheet.column_dimensions[col_letter].width = 20
    sheet.column_dimensions['I'].width = 30
    sheet.column_dimensions['J'].width = 45

    for current_folder, _, files in os.walk(root_folder):
        for file_name in files:
            if file_name.endswith(".docx") and not file_name.startswith("~$"):
                full_path = os.path.join(current_folder, file_name)
                print(f"Processing file: {full_path}")
                extracted_data = extract_word_data(full_path)
                if extracted_data:
                    sheet.append(
                        [
                            extracted_data["product_name"],
                            extracted_data["client"],
                            extracted_data["production_per_hour"],
                            extracted_data["channel_break"],
                            extracted_data["deburring"],
                            extracted_data["sandblasting"],
                            extracted_data["sanding"],
                            extracted_data["packaging_inspection"],
                            extracted_data["packaging_type"],
                            extracted_data["quantity_per_package"],
                        ]
                    )
                    for cell in sheet[sheet.max_row]:
                        cell.alignment = center_align
                else:
                    print(f"File {full_path} skipped due to extraction error.")

    workbook.save(output_excel_file)
    print(f"Excel file generated: {output_excel_file}")

if __name__ == "__main__":
    current_working_directory = os.getcwd()
    output_excel_filename = "produtos_simplificado.xlsx"
    process_folders_and_create_excel(current_working_directory, output_excel_filename)